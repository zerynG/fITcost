from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
from django import forms
from django.db.models import Q
from .models import CommercialProposal
from .forms import CommercialProposalForm, ServiceItemFormSet
from openpyxl import Workbook
from docx import Document
from datetime import datetime
from io import BytesIO
from xhtml2pdf import pisa
from django.urls import reverse

def debug_urls(request):
    """Временная функция для отладки URL"""
    urls_to_test = [
        'commercial_proposal:proposal_list',
        'commercial_proposal:create_proposal',
        'commercial_proposal:proposal_detail',
    ]

    result = "<h1>Debug URLs</h1>"
    for url_name in urls_to_test:
        try:
            url = reverse(url_name)
            result += f"<p>✓ {url_name} -> {url}</p>"
        except Exception as e:
            result += f"<p>✗ {url_name} -> ERROR: {e}</p>"

    return HttpResponse(result)


@login_required
def proposal_list(request, workspace_id=None, project_id=None):
    """Список всех коммерческих предложений с фильтрацией по проекту"""
    from workspace.models import Project
    
    if project_id:
        project = get_object_or_404(Project, id=project_id)
        # Показываем коммерческие предложения только для текущего проекта
        proposals = CommercialProposal.objects.filter(project=project).order_by('-creation_date')
    else:
        project = None
        # Показываем все коммерческие предложения, если проект не указан
        proposals = CommercialProposal.objects.all().order_by('-creation_date')
    
    return render(request, 'commercial_proposal/list.html', {
        'proposals': proposals,
        'project_id': project_id,
        'workspace_id': workspace_id,
        'project': project,
    })

@login_required
def create_proposal(request, workspace_id=None, project_id=None):
    from workspace.models import Project
    from customers.models import Customer
    
    project = Project.objects.filter(id=project_id).first() if project_id else None

    def filter_customers(form_obj):
        if project:
            qs = Customer.objects.filter(
                Q(project=project) | Q(can_be_shared=True, project__workspace=getattr(project, "workspace", None))
            ).distinct()
        else:
            qs = Customer.objects.none()
        if "customer" in form_obj.fields:
            form_obj.fields["customer"].queryset = qs
    
    if request.method == 'POST':
        form = CommercialProposalForm(request.POST)
        filter_customers(form)
        formset = ServiceItemFormSet(request.POST)

        if form.is_valid() and formset.is_valid():
            proposal = form.save(commit=False)
            if project:
                proposal.project = project
            proposal.save()
            formset.instance = proposal
            formset.save()
            if project_id:
                return redirect('commercial_proposal:proposal_list_project', workspace_id=workspace_id, project_id=project_id)
            return redirect('commercial_proposal:proposal_detail', pk=proposal.pk)
    else:
        form = CommercialProposalForm()
        if project:
            form.fields['project'].initial = project
            form.fields['project'].widget = forms.HiddenInput()
        filter_customers(form)
        formset = ServiceItemFormSet()

    context = {
        'form': form,
        'formset': formset,
        'project_id': project_id,
        'workspace_id': workspace_id,
        'project': project,
    }
    return render(request, 'commercial_proposal/create.html', context)


def proposal_detail(request, pk, workspace_id=None, project_id=None):
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    # Если project_id не передан, пытаемся получить из proposal
    if not project_id and proposal.project:
        project_id = proposal.project.id
        if proposal.project.workspace:
            workspace_id = proposal.project.workspace.id
    return render(request, 'commercial_proposal/detail.html', {
        'proposal': proposal,
        'project_id': project_id,
        'workspace_id': workspace_id,
    })


@login_required
def edit_proposal(request, pk, workspace_id=None, project_id=None):
    """Редактирование коммерческого предложения"""
    from workspace.models import Project
    from customers.models import Customer
    
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    # Если project_id не передан, пытаемся получить из proposal
    if not project_id and proposal.project:
        project_id = proposal.project.id
        if proposal.project.workspace:
            workspace_id = proposal.project.workspace.id
    
    project = Project.objects.filter(id=project_id).first() if project_id else None

    def filter_customers(form_obj):
        if project:
            qs = Customer.objects.filter(
                Q(project=project) | Q(can_be_shared=True, project__workspace=getattr(project, "workspace", None))
            ).distinct()
        else:
            qs = Customer.objects.all()
        if "customer" in form_obj.fields:
            form_obj.fields["customer"].queryset = qs
    
    if request.method == 'POST':
        form = CommercialProposalForm(request.POST, instance=proposal)
        filter_customers(form)
        formset = ServiceItemFormSet(request.POST, instance=proposal)

        if form.is_valid() and formset.is_valid():
            proposal = form.save(commit=False)
            if project:
                proposal.project = project
            proposal.save()
            formset.save()
            if project_id and workspace_id:
                return redirect('commercial_proposal:proposal_detail_project', workspace_id=workspace_id, project_id=project_id, pk=proposal.pk)
            return redirect('commercial_proposal:proposal_detail', pk=proposal.pk)
    else:
        form = CommercialProposalForm(instance=proposal)
        if project:
            form.fields['project'].initial = project
            form.fields['project'].widget = forms.HiddenInput()
        filter_customers(form)
        formset = ServiceItemFormSet(instance=proposal)

    context = {
        'form': form,
        'formset': formset,
        'proposal': proposal,
        'project_id': project_id,
        'workspace_id': workspace_id,
        'project': project,
    }
    return render(request, 'commercial_proposal/create.html', context)


def delete_proposal(request, pk, workspace_id=None, project_id=None):
    """Удаление коммерческого предложения с подтверждением."""
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    # Если project_id не передан, пытаемся получить из proposal
    if not project_id and proposal.project:
        project_id = proposal.project.id
        if proposal.project.workspace:
            workspace_id = proposal.project.workspace.id
    if request.method == "POST":
        proposal.delete()
        if project_id and workspace_id:
            return redirect('commercial_proposal:proposal_list_project', workspace_id=workspace_id, project_id=project_id)
        return redirect('commercial_proposal:proposal_list')
    return render(request, 'commercial_proposal/confirm_delete.html', {
        'proposal': proposal,
        'project_id': project_id,
        'workspace_id': workspace_id,
    })


def download_pdf(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    html = render_to_string('commercial_proposal/template_pdf.html', {'proposal': proposal})

    # Создаем PDF
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)

    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.pdf"'
        return response

    return HttpResponse("Ошибка при создании PDF", status=500)


def download_excel(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)

    wb = Workbook()
    ws = wb.active
    ws.title = "Коммерческое предложение"

    # Заголовок
    ws['A1'] = proposal.title
    ws['A2'] = f"Дата формирования: {proposal.creation_date}"
    ws['A3'] = f"Заказчик: {proposal.customer.name if proposal.customer else 'Не указан'}"

    # Услуги
    ws['A5'] = "Услуги"
    ws['A6'] = "Название"
    ws['B6'] = "Часы"
    ws['C6'] = "Период"
    ws['D6'] = "Стоимость"

    for i, service in enumerate(proposal.services.all(), 7):
        ws[f'A{i}'] = service.name
        ws[f'B{i}'] = float(service.hours)
        ws[f'C{i}'] = f"{service.start_date} - {service.end_date}"
        ws[f'D{i}'] = float(service.cost)

    # Итог
    last_row = 7 + len(proposal.services.all())
    ws[f'A{last_row}'] = "ИТОГО:"
    ws[f'D{last_row}'] = float(proposal.total_cost)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.xlsx"'
    wb.save(response)
    return response


def download_word(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)

    doc = Document()
    doc.add_heading(proposal.title, 0)
    doc.add_paragraph(f"Дата формирования: {proposal.creation_date}")
    doc.add_paragraph(f"Заказчик: {proposal.customer.name if proposal.customer else 'Не указан'}")

    doc.add_heading('Техническое задание', level=1)
    doc.add_paragraph(proposal.technical_spec)

    doc.add_heading('Услуги', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название'
    hdr_cells[1].text = 'Часы'
    hdr_cells[2].text = 'Период'
    hdr_cells[3].text = 'Стоимость'

    for service in proposal.services.all():
        row_cells = table.add_row().cells
        row_cells[0].text = service.name
        row_cells[1].text = str(service.hours)
        row_cells[2].text = f"{service.start_date} - {service.end_date}"
        row_cells[3].text = str(service.cost)

    doc.add_paragraph(f"ИТОГО: {proposal.total_cost} руб.")
    doc.add_paragraph(f"{proposal.manager_position} {proposal.manager_name}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.docx"'
    doc.save(response)
    return response